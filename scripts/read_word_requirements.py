#!/usr/bin/env python
"""从 Word 文档读取需求，批量生成测试用例"""

import os
import re
import sys
import argparse
from pathlib import Path
from typing import List, Dict, Any

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from test.ai.test_generator import AITestGenerator


def read_word_docx(file_path: str) -> str:
    """读取 Word 文档内容（需要安装 python-docx）"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    content.append(row_text)
        
        return "\n".join(content)
    except ImportError:
        print("❌ 请先安装 python-docx: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"❌ 读取 Word 文件失败: {e}")
        sys.exit(1)


def parse_requirements_by_markdown(content: str) -> List[Dict[str, str]]:
    """
    按 Markdown 格式解析需求
    支持格式：
    ## 模块名
    - 需求描述
    """
    requirements = []
    current_module = None
    current_feature = None
    
    lines = content.split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 二级标题 ## 模块名
        if line.startswith("## "):
            current_module = line[3:].strip()
            current_feature = current_module
        
        # 三级标题 ### 功能名
        elif line.startswith("### "):
            current_feature = line[4:].strip()
        
        # 列表项 - 需求描述
        elif line.startswith("- "):
            if current_module:
                requirement_text = line[2:].strip()
                requirements.append({
                    "module": current_module.upper(),
                    "feature": current_feature or current_module,
                    "requirement": requirement_text,
                    "priority": extract_priority(requirement_text)
                })
    
    return requirements


def extract_priority(text: str) -> str:
    """从需求文本中提取优先级"""
    if "[P0]" in text or "【P0】" in text or "必须" in text:
        return "P0"
    elif "[P1]" in text or "【P1】" in text or "应该" in text:
        return "P1"
    elif "[P2]" in text or "【P2】" in text or "可以" in text:
        return "P2"
    return "P1"  # 默认


def parse_requirements_by_sections(content: str) -> List[Dict[str, str]]:
    """
    按固定格式解析需求（适用于非 Markdown 格式）
    格式示例：
    【模块】LOGIN
    【功能】登录表单
    【需求】用户可以使用邮箱和密码登录
    【优先级】P0
    ---
    """
    requirements = []
    lines = content.split("\n")
    
    current = {}
    separator = "---"
    
    for line in lines:
        line = line.strip()
        
        if line == separator and current:
            if current.get("module") and current.get("requirement"):
                requirements.append(current)
            current = {}
        
        elif line.startswith("【模块】") or line.startswith("[模块]"):
            current["module"] = line.split("】")[-1].strip() if "】" in line else line.split("]")[-1].strip()
            current["module"] = current["module"].upper()
        
        elif line.startswith("【功能】") or line.startswith("[功能]"):
            current["feature"] = line.split("】")[-1].strip() if "】" in line else line.split("]")[-1].strip()
        
        elif line.startswith("【需求】") or line.startswith("[需求]"):
            current["requirement"] = line.split("】")[-1].strip() if "】" in line else line.split("]")[-1].strip()
        
        elif line.startswith("【优先级】") or line.startswith("[优先级]"):
            current["priority"] = line.split("】")[-1].strip() if "】" in line else line.split("]")[-1].strip()
    
    # 添加最后一个
    if current.get("module") and current.get("requirement"):
        requirements.append(current)
    
    return requirements


def parse_requirements_simple(content: str) -> List[Dict[str, str]]:
    """
    简单解析：每行一个需求，自动推断模块
    格式：模块名: 需求描述
    """
    requirements = []
    
    for line in content.split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        
        # 格式：模块名: 需求描述
        if ":" in line:
            parts = line.split(":", 1)
            module = parts[0].strip().upper()
            requirement_text = parts[1].strip()
        else:
            # 没有模块名，需要用户手动指定
            module = "UNKNOWN"
            requirement_text = line
        
        if requirement_text:
            requirements.append({
                "module": module,
                "feature": requirement_text[:30],  # 取前30字作为功能名
                "requirement": requirement_text,
                "priority": extract_priority(requirement_text)
            })
    
    return requirements


def generate_test_cases(
    requirements: List[Dict[str, str]],
    api_key: str = None,
    output_dir: str = "test/ai/generated"
) -> List[Dict[str, Any]]:
    """批量生成测试用例"""
    generator = AITestGenerator(api_key=api_key)
    results = []
    
    for i, req in enumerate(requirements, 1):
        print(f"\n[{i}/{len(requirements)}] 正在生成: {req['module']} - {req['feature']}")
        
        try:
            result = generator.generate_complete_test_case(
                requirement=req["requirement"],
                module_name=req["module"],
                feature_name=req["feature"],
                priority=req.get("priority", "P1")
            )
            
            # 保存到文件
            module_dir = os.path.join(output_dir, req["module"].lower())
            os.makedirs(module_dir, exist_ok=True)
            
            # 保存测试点
            points_file = os.path.join(module_dir, f"{req['feature']}_test_points.md")
            with open(points_file, "w", encoding="utf-8") as f:
                f.write(result["test_points"])
            
            # 保存测试代码
            code_file = os.path.join(module_dir, f"test_{req['feature'].lower()}.py")
            with open(code_file, "w", encoding="utf-8") as f:
                f.write(result["test_code"])
            
            results.append({
                "module": req["module"],
                "feature": req["feature"],
                "status": "success",
                "files": [points_file, code_file]
            })
            
            print(f"  ✅ 已保存到 {module_dir}/")
            
        except Exception as e:
            print(f"  ❌ 生成失败: {e}")
            results.append({
                "module": req["module"],
                "feature": req["feature"],
                "status": "failed",
                "error": str(e)
            })
    
    return results


def print_summary(results: List[Dict[str, Any]]):
    """打印生成摘要"""
    print("\n" + "=" * 60)
    print("📊 批量生成完成")
    print("=" * 60)
    
    success_count = sum(1 for r in results if r["status"] == "success")
    failed_count = len(results) - success_count
    
    print(f"✅ 成功: {success_count}")
    print(f"❌ 失败: {failed_count}")
    
    if failed_count > 0:
        print("\n失败列表:")
        for r in results:
            if r["status"] == "failed":
                print(f"  - {r['module']}/{r['feature']}: {r.get('error', '未知错误')}")


def main():
    parser = argparse.ArgumentParser(description="从 Word 文档读取需求，批量生成测试用例")
    parser.add_argument("--word", "-w", type=str, help="Word 文档路径 (.docx)")
    parser.add_argument("--text", "-t", type=str, help="文本文件路径 (.txt)")
    parser.add_argument("--format", "-f", type=str, default="auto", 
                        choices=["auto", "markdown", "section", "simple"],
                        help="需求文档格式: auto(自动), markdown, section(分区), simple(简单)")
    parser.add_argument("--output", "-o", type=str, default="test/ai/generated", help="输出目录")
    parser.add_argument("--api-key", type=str, help="DeepSeek API Key")
    
    args = parser.parse_args()
    
    # 读取内容
    if args.word:
        content = read_word_docx(args.word)
    elif args.text:
        with open(args.text, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        # 交互式输入
        print("请输入需求内容（输入空行结束，支持 Markdown 格式）:")
        lines = []
        while True:
            line = input()
            if not line:
                break
            lines.append(line)
        content = "\n".join(lines)
    
    # 解析需求
    if args.format == "auto":
        # 自动判断格式
        if "## " in content:
            requirements = parse_requirements_by_markdown(content)
        elif "【模块】" in content or "[模块]" in content:
            requirements = parse_requirements_by_sections(content)
        else:
            requirements = parse_requirements_simple(content)
    elif args.format == "markdown":
        requirements = parse_requirements_by_markdown(content)
    elif args.format == "section":
        requirements = parse_requirements_by_sections(content)
    else:  # simple
        requirements = parse_requirements_simple(content)
    
    if not requirements:
        print("❌ 未找到有效的需求内容")
        return
    
    print(f"\n📋 解析到 {len(requirements)} 条需求")
    for req in requirements:
        print(f"  - [{req['module']}] {req['requirement'][:50]}...")
    
    # 确认生成
    confirm = input("\n是否开始生成测试用例？(y/N): ").strip().lower()
    if confirm != 'y':
        print("已取消")
        return
    
    # 批量生成
    results = generate_test_cases(requirements, args.api_key, args.output)
    
    # 打印摘要
    print_summary(results)


if __name__ == "__main__":
    main()
#!/usr/bin/env python
"""将 Word 需求文档转换为 Markdown 格式"""

import os
import re
import argparse
from pathlib import Path


def convert_word_to_markdown(docx_path: str, output_path: str = None) -> str:
    """
    将 Word 文档转换为 Markdown 格式
    
    转换规则：
    - 标题1（Heading 1）→ # 标题
    - 标题2（Heading 2）→ ## 标题
    - 标题3（Heading 3）→ ### 标题
    - 列表段落 → - 内容
    - 表格 → Markdown 表格
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
    except ImportError:
        print("❌ 请先安装 python-docx: pip install python-docx")
        return None
    
    doc = Document(docx_path)
    markdown_lines = []
    
    # 添加文档来源信息
    markdown_lines.append(f"<!-- 来源: {os.path.basename(docx_path)} -->")
    markdown_lines.append("")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # 获取段落样式
        style_name = para.style.name.lower() if para.style else ""
        
        # 标题转换
        if "heading 1" in style_name or "标题 1" in style_name:
            markdown_lines.append(f"# {text}")
        elif "heading 2" in style_name or "标题 2" in style_name:
            markdown_lines.append(f"## {text}")
        elif "heading 3" in style_name or "标题 3" in style_name:
            markdown_lines.append(f"### {text}")
        elif "heading 4" in style_name or "标题 4" in style_name:
            markdown_lines.append(f"#### {text}")
        
        # 列表项转换
        elif para._element.p is not None:
            # 检查是否是列表项
            pPr = para._element.p.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                # 获取缩进层级确定列表级别
                ilvl = numPr.find(qn('w:ilvl'))
                level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0
                indent = "  " * level
                markdown_lines.append(f"{indent}- {text}")
            else:
                markdown_lines.append(text)
        else:
            markdown_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        markdown_lines.append("")
        
        # 表头
        header_cells = []
        for cell in table.rows[0].cells:
            header_cells.append(cell.text.strip())
        markdown_lines.append("| " + " | ".join(header_cells) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # 数据行
        for row in table.rows[1:]:
            row_cells = []
            for cell in row.cells:
                row_cells.append(cell.text.strip())
            markdown_lines.append("| " + " | ".join(row_cells) + " |")
        
        markdown_lines.append("")
    
    # 输出 Markdown 内容
    md_content = "\n".join(markdown_lines)
    
    # 保存文件
    if output_path is None:
        output_path = str(Path(docx_path).with_suffix('.md'))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    print(f"✅ 已转换: {docx_path} → {output_path}")
    return md_content


def main():
    parser = argparse.ArgumentParser(description="将 Word 需求文档转换为 Markdown 格式")
    parser.add_argument("input", type=str, help="Word 文档路径 (.docx)")
    parser.add_argument("-o", "--output", type=str, help="输出 Markdown 文件路径（可选）")
    
    args = parser.parse_args()
    
    if not args.input.endswith('.docx'):
        print("❌ 请提供 .docx 文件")
        return
    
    convert_word_to_markdown(args.input, args.output)


if __name__ == "__main__":
    main()